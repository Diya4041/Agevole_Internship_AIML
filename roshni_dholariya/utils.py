# # utils.py
# from pdfminer.high_level import extract_text
# from docx import Document
# import os

# def extract_text_from_file(path):
#     ext = os.path.splitext(path)[-1].lower()
#     if ext == '.pdf':
#         return extract_text(path)
#     elif ext == '.docx':
#         doc = Document(path)
#         return '\n'.join(p.text for p in doc.paragraphs)
#     return ''
# utils.py
import os
from pdfminer.high_level import extract_text
from docx import Document

def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == ".pdf":
        return extract_text(file_path)
    elif ext == ".docx":
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        return ""
